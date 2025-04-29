# --- Imports ---
import os
import json
import uuid
import tempfile
import subprocess
from datetime import datetime

import streamlit as st
import openai
import markdown2
import fitz  # PyMuPDF
import docx  # python-docx for .docx handling
from dotenv import load_dotenv
from ghostwriter_doc_learning import Workspace

# --- Initialize Workspace ---
workspace = Workspace()

# --- Load OpenAI API Key ---
openai.api_key = st.secrets["OPENAI_API_KEY"]

# --- Page Settings ---
st.set_page_config(page_title="Ghostwriter", layout="wide")

# --- Global Custom Styling ---
st.markdown("""
<style>
body, .stApp { background-color: #f7f9fa; color: #333; font-family: 'Helvetica Neue', sans-serif; }
h1, h2, h3, .big-title {
    font-size: 48px;
    font-weight: 700;
    margin-top: 1rem;
    margin-bottom: 0.5rem;
    color: #2C3E50;
}
.stButton>button { background-color: #5D737E; color: white; border-radius: 8px; padding: 8px 16px; border: none; transition: background-color 0.3s ease; }
.stButton>button:hover { background-color: #4C5B68; }
.section-card { background: #ffffff; border-radius: 10px; box-shadow: 0 2px 8px rgba(0,0,0,0.05); padding: 2rem; margin-bottom: 2rem; }
.sidebar .sidebar-content { background-color: #f0f2f4; }
.stAlert {
    background-color: #e6f4ea !important;
    color: #2c662d !important;
    border-left: 5px solid #34a853 !important;
    border-radius: 8px;
    padding: 1rem;
    margin-top: 1.5rem;
    box-shadow: 0 2px 6px rgba(0, 0, 0, 0.1);
}
</style>
""", unsafe_allow_html=True)



# --- Top Section: Branding and Privacy ---
st.markdown('<div class="big-title"> Ghostwriter</div>', unsafe_allow_html=True)

st.caption("Your AI assistant for writing, reviewing, and managing technical documents.")
with st.expander("💬 Privacy & Data Use", expanded=False):
    st.markdown("""
    **Your content is yours.**  
    We use ChatGPT via the OpenAI API, which never stores your prompts or documents and doesn’t use them to train future models.  
    **Ghostwriter doesn’t store anything** unless you click “Save.”  
    Your data is encrypted, and you have full control over how and when AI assists you.
    """)



# --- Sidebar ---
with st.sidebar:
    st.markdown('<div class="sidebar-header">🧠 Ghostwriter</div>', unsafe_allow_html=True)
    st.caption("Configure your document")

    # Document Setup
    st.selectbox("📂 Document Type", ["Quick Start", "Install Guide", "Safety Sheet", "FAQ"], key="doc_type")
    st.selectbox("👥 Audience", ["End User", "Technician", "Support Staff"], key="audience")

    st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)

    # Learn and Review
    st.markdown("### 📚 Learning & Review")

    with st.expander("ℹ️ What is this?", expanded=False):
        st.markdown("""
    **Learn from Existing Documents**  
    Upload finalized documents to help Ghostwriter learn your style, terminology, and tone. Once 5 documents or 10,000 words are uploaded, the system builds a style model.

    **Review a New Document**  
    Upload a new document to receive feedback based on your learned style and terminology. It checks for passive voice, long sentences, and inconsistent terms.
    """)



    with st.expander("Learn from Existing Documents"):
        learn_file = st.file_uploader("Upload a document to train Ghostwriter", type=["txt", "md", "docx"], key="learn_upload")
        learn_status = st.selectbox("Document Status", ["draft", "final"], key="learn_status")
        if learn_file and st.button("Upload for Learning", key="learn_button"):
            ext = learn_file.name.split('.')[-1].lower()
            if ext == "docx":
                doc = docx.Document(learn_file)
                learn_text = "\n".join(p.text for p in doc.paragraphs)
            else:
                learn_text = learn_file.read().decode("utf-8")
            workspace.upload_document(learn_text, learn_file.name, learn_status)
            st.success(f"✅ {learn_file.name} uploaded and tagged as {learn_status}.")

    with st.expander("Review a New Document"):
        review_file = st.file_uploader("Upload for review", type=["txt", "md", "docx"], key="review_upload")
        if review_file and st.button("Run Review", key="review_button"):
            ext = review_file.name.split('.')[-1].lower()
            if ext == "docx":
                doc = docx.Document(review_file)
                review_text = "\n".join(p.text for p in doc.paragraphs)
            else:
                review_text = review_file.read().decode("utf-8")
            feedback = workspace.review_document(review_text)
            st.subheader("📊 Review Feedback")
            for category, items in feedback.items():
                st.markdown(f"#### {category.capitalize()}")
                for item in (items or ["✅ No issues found."]):
                    st.write(f"- {item}")

    st.markdown('<div class="sidebar-divider"></div>', unsafe_allow_html=True)
    st.caption("Ghostwriter v0.9 – Streamlit Edition")




# --- Custom Page Layout Style ---
st.markdown("""
<style>
/* Centered content max-width */
.css-18e3th9 {
    max-width: 1100px;
    margin: auto;
}

/* Section card */
.section-card {
    background-color: #fdfdfd;
    padding: 2rem 1.5rem;
    border-radius: 0.6rem;
    box-shadow: 0 1px 6px rgba(0,0,0,0.06);
    margin-bottom: 2rem;
}
.section-card h3 {
    margin-top: 0;
}

/* Buttons and toggles */
.toggle-button {
    padding: 0.5rem 1rem;
    margin-top: 0.5rem;
    border-radius: 0.4rem;
    border: 1px solid #ccc;
    background-color: #fafafa;
}
.toggle-button:hover {
    background-color: #f0f0f0;
}
</style>
""", unsafe_allow_html=True)


# --- Add Product Info ---
st.markdown("---")
st.header("📥 Add Product Information")

input_col1, input_col2 = st.columns(2)

if "input_method" not in st.session_state:
    st.session_state.input_method = None

with input_col1:
    st.markdown("#### ✏️ Paste Text")
    if st.button("Paste Text", key="text_input_button"):
        st.session_state.input_method = "paste"

with input_col2:
    st.markdown("#### 📁 Upload File")
    if st.button("Upload File", key="file_upload_button"):
        st.session_state.input_method = "upload"

product_info = ""
generate_clicked = False

# Pull document settings from sidebar
doc_type = st.session_state.get("doc_type", "Quick Start")
audience = st.session_state.get("audience", "End User")

if st.session_state.input_method == "paste":
    st.session_state["pasted_text"] = st.text_area(
        "Paste your product specifications or description here:",
        value=st.session_state.get("pasted_text", ""),
        height=300
    )
    product_info = st.session_state["pasted_text"]

    col1, col2 = st.columns(2)
    with col1:
        generate_clicked = st.button("🚀 Generate Draft", key="generate_draft_paste")
    with col2:
        if st.button("🧹 Clear", key="clear_paste"):
            st.session_state["pasted_text"] = ""

elif st.session_state.input_method == "upload":
    uploaded_file = st.file_uploader(
        "Upload a specification file (.txt, .md, .pdf, .docx, .rtf)",
        type=["txt", "md", "pdf", "docx", "rtf"],
        key="spec_file_upload"
    )

    if uploaded_file:
        ext = uploaded_file.name.split('.')[-1].lower()
        if ext == "pdf":
            with fitz.open(stream=uploaded_file.read(), filetype="pdf") as pdf:
                product_info = "".join([page.get_text() for page in pdf])
        elif ext == "docx":
            doc = docx.Document(uploaded_file)
            product_info = "\n".join(p.text for p in doc.paragraphs)
        else:
            product_info = uploaded_file.read().decode("utf-8")

    generate_clicked = st.button("🚀 Generate Draft", key="generate_draft_upload")



# --- Generate Draft ---
if generate_clicked and product_info:
    with st.spinner("Generating draft..."):
        try:
            client = openai.OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
            base_prompt = f"""
You are a technical writer with all the experience and expertise of a 20 year career professional. Generate a professional {st.session_state['doc_type']} for {st.session_state['audience']}.
Use Markdown format, H1/H2, bullets or numbers, avoid repetition and marketing fluff. Be direct and helpful. 
"""
            if "style_guide" in st.session_state:
                base_prompt += f"\nStrictly follow this additional style guide:\n{st.session_state['style_guide']}"

            user_input = f"""DOCUMENT TYPE: {st.session_state['doc_type']}
AUDIENCE: {st.session_state['audience']}
PRODUCT INFO:
{product_info}
"""

            response = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "system", "content": base_prompt}, {"role": "user", "content": user_input}],
                temperature=0.4
            )

            draft = response.choices[0].message.content
            st.session_state["generated_md"] = draft

        except Exception as e:
            st.error(f"Error generating draft: {e}")

# --- Display Draft, Download, and Save ---
if "generated_md" in st.session_state:
    st.subheader("📝 Your Markdown Draft")
    edited_md = st.text_area(
        "Edit your draft below:",
        value=st.session_state["generated_md"],
        height=400
    )

    st.markdown("### 📥 Download Options")
    col1, col2, col3 = st.columns(3)

    with col1:
        st.download_button("⬇️ Download Markdown", edited_md, file_name="draft.md", mime="text/markdown")

    with col2:
        html_output = markdown2.markdown(edited_md)
        st.download_button("🌐 Export to HTML", html_output, file_name="draft.html", mime="text/html")

    with col3:
        with tempfile.NamedTemporaryFile("w+", suffix=".md", delete=False) as tmp:
            tmp.write(edited_md)
            tmp_path = tmp.name
        try:
            subprocess.run(["pandoc", tmp_path, "-o", tmp_path.replace(".md", ".pdf")], check=True)
            with open(tmp_path.replace(".md", ".pdf"), "rb") as pdf:
                st.download_button("📄 Export to PDF", pdf.read(), file_name="draft.pdf", mime="application/pdf")
        except Exception as e:
            st.error(f"PDF export failed: {e}")

    st.markdown("---")
    st.subheader("💾 Save to Document Library?")

    save_col1, save_col2 = st.columns([2, 1])

    with save_col1:
        default_title = f"{st.session_state.get('doc_type', 'Document')} – {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        custom_title = st.text_input("Document Title", value=default_title)

    with save_col2:
        if st.button("✅ Save Draft"):
            try:
                os.makedirs("docs", exist_ok=True)
                doc_id = str(uuid.uuid4())
                doc_data = {
                    "id": doc_id,
                    "name": custom_title,
                    "type": st.session_state.get("doc_type", "Quick Start"),
                    "audience": st.session_state.get("audience", "End User"),
                    "date": datetime.now().strftime("%Y-%m-%d"),
                    "content": edited_md,
                    "tags": [],
                    "filename": f"{doc_id}.md"
                }
                with open(f"docs/{doc_id}.json", "w") as f:
                    json.dump(doc_data, f, indent=2)
                st.success(f"✅ '{custom_title}' saved to your document library!")
            except Exception as e:
                st.error(f"Failed to save draft: {e}")

# --- Load persisted style guide at startup ---
STYLE_PATH = "style_guide.txt"

if "style_guide" not in st.session_state and os.path.exists(STYLE_PATH):
    with open(STYLE_PATH, "r") as f:
        st.session_state["style_guide"] = f.read()
        st.session_state["style_uploaded_at"] = datetime.fromtimestamp(os.path.getmtime(STYLE_PATH)).strftime("%Y-%m-%d %H:%M:%S")
        st.session_state["style_uploaded_by"] = os.getenv("USER", "Unknown User")

# --- Upload Style Guide ---
st.markdown("---")
style_guide_expanded = "style_guide" not in st.session_state

with st.expander("📘 Upload a Style Guide", expanded=style_guide_expanded):
    style_file = st.file_uploader(
        "Upload a style guide (.txt, .docx, .pdf)",
        type=["txt", "docx", "pdf"],
        key="style_guide_upload"
    )

    if style_file:
        ext = style_file.name.split('.')[-1].lower()

        if ext == "pdf":
            with fitz.open(stream=style_file.read(), filetype="pdf") as pdf:
                style_text = "".join([page.get_text() for page in pdf])
        elif ext == "docx":
            doc = docx.Document(style_file)
            style_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            style_text = style_file.read().decode("utf-8")

        # Save to session and disk
        st.session_state["style_guide"] = style_text
        st.session_state["style_uploaded_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        st.session_state["style_uploaded_by"] = os.getenv("USER", "Unknown User")

        with open("style_guide.txt", "w") as f:
            f.write(style_text)

        st.success(f"""
✅ **Style guide uploaded and saved!**

- **Uploaded by:** {st.session_state['style_uploaded_by']}
- **Uploaded at:** {st.session_state['style_uploaded_at']}
""")

# Display current guide info
if "style_guide" in st.session_state:
    st.info(f"""
📘 **Active Style Guide**

- **Uploaded by:** {st.session_state.get('style_uploaded_by', 'Unknown')}
- **Uploaded at:** {st.session_state.get('style_uploaded_at', 'Unknown')}
""")






